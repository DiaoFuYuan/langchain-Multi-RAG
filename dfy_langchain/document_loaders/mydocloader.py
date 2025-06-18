from typing import List
import os
import tqdm
from langchain_community.document_loaders.unstructured import UnstructuredFileLoader
from PIL import Image
import numpy as np
from io import BytesIO


class RapidOCRDocLoader(UnstructuredFileLoader):
    """
    适用于Word文档(.doc和.docx)的加载器，具有OCR功能，可以提取文档中的文本和图像内容。
    
    使用说明:
    - 对于.docx文件，使用python-docx库提取文本和图像，并使用RapidOCR进行图像OCR
    - 对于.doc文件，使用unstructured库直接处理
    """
    
    def _get_elements(self) -> List:
        # 检查文件格式
        file_ext = os.path.splitext(self.file_path)[1].lower()
        
        if file_ext == '.docx':
            # 对于.docx文件，使用我们的自定义处理以支持OCR
            text = self._process_docx(self.file_path)
            from unstructured.partition.text import partition_text
            return partition_text(text=text, **self.unstructured_kwargs)
        elif file_ext == '.doc':
            # 对于.doc文件，直接使用unstructured库
            from unstructured.partition.doc import partition_doc
            return partition_doc(filename=self.file_path, **self.unstructured_kwargs)
        else:
            raise ValueError(f"不支持的文档格式: {file_ext}，仅支持.doc和.docx格式")
    
    def _process_docx(self, filepath):
        """处理.docx格式文件，支持文本提取和图像OCR"""
        from docx import Document, ImagePart
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table, _Cell
        from docx.text.paragraph import Paragraph
        from rapidocr_onnxruntime import RapidOCR

        ocr = RapidOCR()
        doc = Document(filepath)
        resp = ""

        def iter_block_items(parent):
            from docx.document import Document

            if isinstance(parent, Document):
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            else:
                raise ValueError("RapidOCRDocLoader parse fail")

            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        b_unit = tqdm.tqdm(
            total=len(doc.paragraphs) + len(doc.tables),
            desc="RapidOCRDocLoader block index: 0",
        )
        for i, block in enumerate(iter_block_items(doc)):
            b_unit.set_description("RapidOCRDocLoader block index: {}".format(i))
            b_unit.refresh()
            if isinstance(block, Paragraph):
                resp += block.text.strip() + "\n"
                images = block._element.xpath(".//pic:pic")  # 获取所有图片
                for image in images:
                    for img_id in image.xpath(".//a:blip/@r:embed"):  # 获取图片id
                        part = doc.part.related_parts[
                            img_id
                        ]  # 根据图片id获取对应的图片
                        if isinstance(part, ImagePart):
                            image = Image.open(BytesIO(part._blob))
                            result, _ = ocr(np.array(image))
                            if result:
                                ocr_result = [line[1] for line in result]
                                resp += "\n".join(ocr_result)
            elif isinstance(block, Table):
                for row in block.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            resp += paragraph.text.strip() + "\n"
            b_unit.update(1)
        return resp


if __name__ == "__main__":
    # 测试.docx文件
    loader = RapidOCRDocLoader(file_path="../tests/samples/ocr_test.docx")
    docs = loader.load()
    print("DOCX 文件测试结果:", docs)
    
    # 测试.doc文件
    try:
        loader = RapidOCRDocLoader(file_path="../tests/samples/test.doc")
        docs = loader.load()
        print("DOC 文件测试结果:", docs)
    except Exception as e:
        print(f"DOC 文件测试失败: {str(e)}")
